from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Request, Body, Query, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import sqlite3
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from typing import List, Dict, Optional, Any
from database.chat_history import ChatHistoryHandler
from processors.code_executor import CodeExecutor
from pydantic import BaseModel
from fastapi.responses import StreamingResponse
import os
import json
import uuid
from datetime import datetime

import hashlib

from models.llm_handler import LlamaModel
from models.embedding import LocalEmbedder
from database.chromadb_handler import ChromaDBHandler
from processors.document_processor import DocumentProcessor
from processors.text_processor import TextProcessor
from utils.config import Settings

settings = Settings()
app = FastAPI(title="Advanced Chatbot API")
# Load app settings from JSON file
# def load_app_settings():
#     try:
#         # Adjust the path based on where you place the JSON file
#         settings_path = os.path.join(os.path.dirname(__file__), "appsettings.json")
#         with open(settings_path, "r") as f:
#             prompts_data = json.load(f)
#         return prompts_data
#     except Exception as e:
#         print(f"Error loading app settings: {e}")
#         # Return a default setting if file can't be loaded
#         return {"default":{ "You are a helpful assistant."
#                              "enable_docs": False}}

def load_app_settings():

    try:

        # Get absolute path to settings file

        settings_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "appsettings.json")

        # Check if file exists

        if not os.path.exists(settings_path):

            print(f"Settings file not found at: {settings_path}")

            return {"default": "You are a helpful assistant.", "enable_docs": False}

        # Read and parse JSON

        with open(settings_path, "r", encoding='utf-8') as f:

            settings = json.load(f)

        print(f"Settings loaded successfully from: {settings_path}")

        return settings

    except json.JSONDecodeError as e:

        print(f"Invalid JSON in settings file: {e}")

        return {"default": "You are a helpful assistant.", "enable_docs": False}

    except Exception as e:

        print(f"Error loading app settings: {e}")

        return {"default": "You are a helpful assistant.", "enable_docs": False}
 

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with specific origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Add this model class with your other models
class CodeExecutionRequest(BaseModel):    
    code: str

# Initialize the code executor
code_executor = CodeExecutor(    
    timeout=settings.CODE_EXECUTION_TIMEOUT,    
    max_output_size=settings.MAX_CODE_OUTPUT_SIZE)

# Get the absolute path to the frontend directory
frontend_dir = os.path.abspath(os.path.join(
    os.path.dirname(__file__),
    '..', 'frontend'
))

# Add these models to your existing models
class MessageCreate(BaseModel):
    role: str
    content: str

class ConversationCreate(BaseModel):
    title: str = "New Chat"
    metadata: Optional[Dict[str, Any]] = None

class ConversationUpdate(BaseModel):
    title: str

# Initialize chat history handler with your app
chat_history = ChatHistoryHandler(db_path="./database/chat_history.db")

# Create a custom StaticFiles class with JavaScript MIME type
class JavaScriptStaticFiles(StaticFiles):
    async def get_response(self, path, scope):
        response = await super().get_response(path, scope)
        if path.endswith('.js'):
            response.headers['Content-Type'] = 'application/javascript'
        return response

class ChatRequest(BaseModel):
    query: str
    context: Optional[str] = None
    conversation_history: Optional[List[Dict[str, str]]] = None

async def serialize_to_sse(data: Dict) -> str:
    """Serialize data to Server-Sent Events format"""
    return f"data: {json.dumps(data)}\n\n"

# Mount static files (CSS, JS, assets)
app.mount("/css", StaticFiles(directory=os.path.join(frontend_dir, "css")), name="css")
app.mount("/js", StaticFiles(directory=os.path.join(frontend_dir, "js")), name="js")
app.mount("/assets", StaticFiles(directory=os.path.join(frontend_dir, "assets")), name="assets")

# Set up templates
templates = Jinja2Templates(directory=frontend_dir)

# @app.get("/")
# async def get_ui(request: Request):
#     # Get all query parameters
#     query_params = dict(request.query_params)
    
#     # Print query parameters to console
#     if query_params:
#         print("URL Parameters:")
#         for key, value in query_params.items():
#             print(f"  {key}: {value}")
    
#     # Pass query parameters to the template
#     context = {"request": request}
#     # If app parameter is present, get config for that app
#     if "app" in query_params:
#         context["app_name"] = query_params["app"]
#         app_name = query_params["app"]

#         # Load prompts from JSON file
#         app_prompts = load_app_settings()
        
#         # Get the prompt for this app, or use default if not found
#         system_prompt = app_prompts.get(app_name, app_prompts.get("default", "You are a helpful assistant."))
#         context["server_system_prompt"] = system_prompt
#         print(f"Using system prompt for {app_name}: {system_prompt}")
#         print(f"Passing app name: {query_params['app']} to template")
    
#     return templates.TemplateResponse("index.html", context)

@app.get("/")
async def get_ui(request: Request):
    # Get all query parameters
    query_params = dict(request.query_params)
    # Print query parameters to console
    if query_params:
        print("URL Parameters:")
        for key, value in query_params.items():
            print(f"  {key}: {value}")
    # Pass query parameters to the template
    context = {"request": request}
    # If app parameter is present, get config for that app
    if "app" in query_params:
        context["app_name"] = query_params["app"]
        app_name = query_params["app"]
 
        try:
            # Load prompts from JSON file
            app_prompts = load_app_settings()
            # Handle nested structure properly
            if app_name in app_prompts:
                app_config = app_prompts[app_name]
                # Check if it's a nested object or direct string
                if isinstance(app_config, dict):
                    system_prompt = app_config.get("system_prompt", "You are a helpful assistant.")
                    enable_docs = app_config.get("enable_docs", False)
                else:
                    # Handle old format (direct string)
                    system_prompt = app_config
                    enable_docs = False
            else:
                # Use default if app not found
                default_config = app_prompts.get("default", {})
                if isinstance(default_config, dict):
                    system_prompt = default_config.get("system_prompt", "You are a helpful assistant.")
                    enable_docs = default_config.get("enable_docs", False)
                else:
                    system_prompt = "You are a helpful assistant."
                    enable_docs = False
            context["server_system_prompt"] = system_prompt
            context["enable_docs"] = enable_docs
            print(f"Using system prompt for {app_name}: {system_prompt[:50]}...")
            print(f"Enable docs: {enable_docs}")
        except Exception as e:
            print(f"Error loading settings for app {app_name}: {e}")
            context["server_system_prompt"] = "You are a helpful assistant."
            context["enable_docs"] = False
    return templates.TemplateResponse("index.html", context)

# Initialize components
embedder = LocalEmbedder()
chroma_db = ChromaDBHandler(embedder)
llm = LlamaModel()
doc_processor = DocumentProcessor()
text_processor = TextProcessor()

# Store chat histories in memory (in production, use a proper database)
chat_histories = {}
system_prompts = {}

# @app.get("/")
# async def root():
#     return {"message": "Chatbot API is running"}

# @app.post("/chat")
# async def chat(
#     message: str = Form(...),
#     chat_id: Optional[str] = Form(None),
#     use_docs: bool = Form(False),
#     system_prompt: Optional[str] = Form(None),
#     # temperature: Optional[float] = Form(None),
#     # max_tokens: Optional[int] = Form(None)
# ):
#     # Create a new chat if chat_id is not provided
#     if not chat_id:
#         chat_id = str(uuid.uuid4())
#         chat_histories[chat_id] = []
#     elif chat_id not in chat_histories:
#         chat_histories[chat_id] = []
    
#     # Add the user message to the history
#     chat_histories[chat_id].append({
#         "role": "user",
#         "content": message,
#         "timestamp": datetime.now().isoformat()
#     })
    
#     # Get or set system prompt
#     if system_prompt:
#         system_prompts[chat_id] = system_prompt
#     current_system_prompt = system_prompts.get(chat_id, "You are a helpful assistant.")
    

#     # Generate response
#     if use_docs:
#         print('using docs')
#         # RAG flow: Query the vector store
#         query_embeddings = embedder.embed_query(message)
#         relevant_docs = chroma_db.similarity_search(query_embeddings)
        
#         # Augment the prompt with the retrieved documents
#         context = "\n\n".join([doc["content"] for doc in relevant_docs])
#         augmented_prompt = f"{current_system_prompt}\n\nContext information:\n{context}\n\nUser question: {message}"
        
#         # Generate response using the LLM
#         response = llm.generate(augmented_prompt, chat_histories[chat_id])
#     else:
#         print('without using docs')
#         # Standard chat flow
#         response = llm.generate(message, chat_histories[chat_id], system_prompt=current_system_prompt)
    
#     # Add the assistant response to the history
#     chat_histories[chat_id].append({
#         "role": "assistant",
#         "content": response,
#         "timestamp": datetime.now().isoformat()
#     })
#     print("HERE--------------> response")
#     return {
#         "response": response,
#         "chat_id": chat_id
#     }

# @app.post("/chat")
# async def chat(
#     message: str = Form(...),
#     chat_id: Optional[str] = Form(None),
#     use_docs: bool = Form(False),
#     system_prompt: Optional[str] = Form(None),
# ):
#     print('[DEBUG]    system_prompt '+system_prompt)
#     if not chat_id:
#         # Create a new conversation in the database
#         metadata = {"system_prompt": system_prompt} if system_prompt else {}
#         chat_id = chat_history.create_conversation(
#             title=message[:30] + "...",  # Use first 30 chars of message as title
#             metadata=metadata
#         )
#     print('Add the user message to the database '+message)
#     try:
#         # Add the user message to the database
#         chat_history.add_message(
#             conversation_id=chat_id,
#             role="user",
#             content=message
#         )
#         print('Message add in chat_history ' + message)  # THIS IS NOT PRINTING
#     except Exception as e:
#         print("Error while adding message to chat_history:", str(e))
#         raise HTTPException(status_code=500, detail=str(e))
    
#     print('Message add in chat_history '+message)
    
#     try:
#         # Get the conversation to retrieve all previous messages
#         print('Getting add in chat_history '+chat_id)
#         conversation = chat_history.get_conversation(chat_id)
#         conversation_history = conversation.get("messages", [])
        
#         # Extract system prompt from metadata if available
#         metadata = conversation.get("metadata", {})
#         current_system_prompt = metadata.get("system_prompt", "You are a helpful assistant.")
        
#         # Update system prompt if provided in this request
#         if system_prompt:
#             # Update metadata to store the new system prompt
#             metadata["system_prompt"] = system_prompt
#             current_system_prompt = system_prompt
            
#             # Update the conversation metadata
#             conn = sqlite3.connect(chat_history.db_path, timeout=10)
#             cursor = conn.cursor()
#             cursor.execute(
#                 "UPDATE conversations SET metadata = ? WHERE conversation_id = ?",
#                 (json.dumps(metadata), chat_id)
#             )
#             conn.commit()
#             conn.close()

#         # Format conversation history for the LLM
#         formatted_history = [
#             {"role": msg["role"], "content": msg["content"]} 
#             for msg in conversation_history
#         ]
#         if use_docs:
#             print('using docs')
#             query_embeddings = embedder.embed_query(message)
#             relevant_docs = chroma_db.similarity_search(query_embeddings)
#             context = "\n\n".join([doc["content"] for doc in relevant_docs])
#             print(f"[DEBUG] Retrieved context: {context[:300]}...")

#             response = ""
#             async for update in llm.generate_response(
#                 query=message,
#                 context=context,
#                 conversation_history=formatted_history,
#                 system_prompt=current_system_prompt
#             ):
#                 if update["status"] == "complete":
#                     response = update["message"]
#         else:
#             print('without using docs')
#             # Standard chat flow
#             response = llm.generate(message, formatted_history, system_prompt=current_system_prompt)
        
#         # Add the assistant response to the database
#         chat_history.add_message(
#             conversation_id=chat_id,
#             role="assistant",
#             content=response
#         )
        
#         return {
#             "response": response,
#             "chat_id": chat_id
#         }
#     except Exception as e:
#         # If there's an error, we should still record it in the database
#         error_message = f"Error generating response: {str(e)}"
#         print('ERROR in Chat ->>>>>'+error_message)
#         chat_history.add_message(
#             conversation_id=chat_id,
#             role="system",
#             content=error_message
#         )
#         raise HTTPException(status_code=500, detail=error_message)

@app.post("/chat")
async def chat(
    message: str = Form(...),
    chat_id: Optional[str] = Form(None),
    use_docs: bool = Form(False),
    system_prompt: Optional[str] = Form(None),
):
    try:
        # Create a new conversation if chat_id is not provided
        print(f'[DEBUG] system_prompt: {system_prompt}')
        if not chat_id:
            # Create a new conversation in the database
            metadata = {"system_prompt": system_prompt} if system_prompt else {}
            chat_id = chat_history.create_conversation(
                title=message[:30] + "...",
                metadata=metadata
            )
        
        print(f'Adding user message to chat_history: {message}')
        
        # Add the user message to the database
        chat_history.add_message(
            conversation_id=chat_id,
            role="user",
            content=message
        )
        
        # Get the conversation to retrieve all previous messages
        conversation = chat_history.get_conversation(chat_id)
        conversation_history = conversation.get("messages", [])
        
        # Extract system prompt from metadata if available
        metadata = conversation.get("metadata", {})
        current_system_prompt = metadata.get("system_prompt", "You are a helpful assistant.")
        
        # Update system prompt if provided in this request
        if system_prompt:
            metadata["system_prompt"] = system_prompt
            current_system_prompt = system_prompt
            
            # Update the conversation metadata
            conn = sqlite3.connect(chat_history.db_path, timeout=300)  # Increased timeout
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE conversations SET metadata = ? WHERE conversation_id = ?",
                (json.dumps(metadata), chat_id)
            )
            conn.commit()
            conn.close()

        # Format conversation history for the LLM
        formatted_history = [
            {"role": msg["role"], "content": msg["content"]} 
            for msg in conversation_history
        ]
        
        # Generate response - Make both paths consistent
        response = ""
        
        if use_docs:
            print('[DEBUG] Using docs for response generation')
            query_embeddings = embedder.embed_query(message)
            relevant_docs = chroma_db.similarity_search(query_embeddings)
            context = "\n\n".join([doc["content"] for doc in relevant_docs])
            print(f"[DEBUG] Retrieved context: {context[:300]}...")

            # Check if generate_response is truly async or sync
            try:
                # Try async approach first
                async for update in llm.generate_response(
                    query=message,
                    context=context,
                    conversation_history=formatted_history,
                    system_prompt=current_system_prompt
                ):
                    if update["status"] == "complete":
                        response = update["message"]
                        break
            except TypeError as e:
                # If async doesn't work, try sync approach
                print(f"[DEBUG] Async generation failed, trying sync: {e}")
                augmented_prompt = f"{current_system_prompt}\n\nContext information:\n{context}\n\nUser question: {message}"
                response = llm.generate(augmented_prompt, formatted_history)
        else:
            print('[DEBUG] Using standard chat without docs')
            # Ensure consistent timeout handling
            try:
                # Try to use async version if available
                if hasattr(llm, 'generate_async'):
                    response = await llm.generate_async(
                        message, 
                        formatted_history, 
                        system_prompt=current_system_prompt
                    )
                else:
                    # Use synchronous version
                    response = llm.generate(
                        message, 
                        formatted_history, 
                        system_prompt=current_system_prompt
                    )
            except Exception as llm_error:
                print(f"[ERROR] LLM Generation failed: {llm_error}")
                response = "I apologize, but I encountered an error while generating a response. Please try again."
        
        # Ensure we got a response
        if not response or response.strip() == "":
            response = "I apologize, but I couldn't generate a proper response. Please try rephrasing your question."
        
        # Add the assistant response to the database
        chat_history.add_message(
            conversation_id=chat_id,
            role="assistant",
            content=response
        )
        
        print(f"[DEBUG] Response generated successfully: {response[:100]}...")
        
        return {
            "response": response,
            "chat_id": chat_id
        }
        
    except sqlite3.OperationalError as db_error:
        print(f'[ERROR] Database error: {db_error}')
        error_message = f"Database timeout error: {str(db_error)}"
        raise HTTPException(status_code=500, detail=error_message)
        
    except Exception as e:
        print(f'[ERROR] General error in /chat: {str(e)}')
        # Try to record the error in the database if possible
        try:
            if 'chat_id' in locals() and chat_id:
                error_message = f"Error generating response: {str(e)}"
                chat_history.add_message(
                    conversation_id=chat_id,
                    role="system",
                    content=error_message
                )
        except:
            pass  # Don't let error recording cause another error
        
        raise HTTPException(status_code=500, detail=f"Error generating response: {str(e)}")

# Create uploads directory if it doesn't exist
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

@app.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = None
):
    
    # Generate a unique filename to avoid overwrites
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_filename = f"{timestamp}_{file.filename}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)

    # Save the file temporarily
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    try:
        # Calculate the hash of the document to check if it has been processed before
        doc_hash = hashlib.sha256(content).hexdigest()
        
        # Check if the document hash is already in the ChromaDB
        if chroma_db.document_exists(doc_hash):
            return {"message": f"Document '{file.filename}' has already been processed."}
        
        # Process the document based on file type
        print('text_content')
        print('JT--- TODO --- YOU CHANGED THE CHUNKING HERE....CHECK TO Continue')
        
        text_content = doc_processor.process_document(file_path)
        print(text_content)

        # Chunk the text (if needed)
        # chunks = text_processor.chunk_text(text_content)
        chunks = doc_processor.chunk_by_heading(text_content)
        
        # Create document IDs for each chunk
        doc_ids = [f"doc_{uuid.uuid4()}" for _ in chunks]
        
        # Embed and store in ChromaDB
        embeddings = embedder.embed_documents(chunks)
        print("DEBUG insides post upload"+file.filename+""+file_path+""+doc_hash)
        # Store the embeddings with the document hash as metadata to track processed docs
        chroma_db.add_documents(doc_ids, chunks, embeddings, {
            "source": file.filename,
            "file_path": file_path,  # Store the path for future reference
            "doc_hash": doc_hash  # Add hash as metadata
        })
        print("DEBUG 2 insides post upload"+file.filename+""+file_path+""+doc_hash)
        return {"message": f"Document '{file.filename}' processed and stored successfully", "file_path": file_path}

    except Exception as e:
        # In case of error, we should remove the file
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")
# @app.get("/documents")
# async def get_documents():
#     # Your logic to retrieve and return documents
#     return {"message": "List of documents"}

class SettingsModel(BaseModel):
    """Model for app settings"""
    system_prompt: Optional[str] = "You are a helpful assistant that provides accurate and concise answers."
    save_chat_history: Optional[bool] = True
    use_docs_toggle: Optional[bool] = False
    chunk_size: Optional[int] = 1000
    overlap_size: Optional[int] = 200
    model_temperature: Optional[float] = 0.7
    model_max_tokens: Optional[int] = 1000
    code_timeout: Optional[int] = 5
    auto_run_code: Optional[bool] = True
    theme_selector: Optional[str] = "system"
    auto_scroll: Optional[bool] = True
    show_thinking_indicator: Optional[bool] = True

# Default settings
default_settings = SettingsModel().dict()

# In-memory settings storage (replace with database in production)
app_settings = default_settings.copy()

@app.get("/settings")
async def get_settings():
    """Get current application settings"""
    return app_settings

@app.post("/settings")
async def save_settings(request: Request):
    """Save application settings"""
    try:
        # Get the raw body as JSON
        data = await request.json()
        
        # Convert kebab-case keys from frontend to snake_case for backend
        for key, value in data.items():
            snake_key = key.replace('-', '_')
            
            # Only update keys that exist in our settings model
            if snake_key in app_settings:
                # Ensure proper type conversion
                if isinstance(app_settings[snake_key], bool):
                    app_settings[snake_key] = bool(value)
                elif isinstance(app_settings[snake_key], int):
                    app_settings[snake_key] = int(value)
                elif isinstance(app_settings[snake_key], float):
                    app_settings[snake_key] = float(value)
                else:
                    app_settings[snake_key] = value
        
        return {"message": "Settings saved successfully", "settings": app_settings}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error saving settings: {str(e)}")    

# @app.delete('/api/documents/delete-all')
# def delete_all_documents():
#     try:
#         # 1. Delete all files in the uploads directory
#         upload_dir = os.path.join(os.path.dirname(__file__), '..', 'uploads')
#         for filename in os.listdir(upload_dir):
#             file_path = os.path.join(upload_dir, filename)
#             if os.path.isfile(file_path):
#                 os.remove(file_path)
        
#         # 2. Clear the ChromaDB collection
#         # Assuming you've instantiated your ChromaDBHandler in app.py
#         # If not, you'll need to import and instantiate it here
#         result = ChromaDBHandler.clear_collection()
#         if not result.get("success", False):
#             return JSONResponse({"success": False, "message": f"Failed to clear ChromaDB: {result.get('error', 'Unknown error')}"}), 500
        
#         # 3. Clear the uploaded_files table in the database
#         db_path = os.path.join(os.path.dirname(__file__), '..', 'database', 'chat_history.db')
#         chat_history = ChatHistoryHandler(db_path)
        
#         # Use SQLite directly for this operation since it's not part of the interface
#         conn = sqlite3.connect(db_path)
#         cursor = conn.cursor()
#         cursor.execute("DELETE FROM uploaded_files")
#         conn.commit()
#         conn.close()
        
#         return JSONResponse({"success": True, "message": "All documents deleted successfully"})
#     except Exception as e:
#         return JSONResponse({"success": False, "message": str(e)}), 500

@app.delete('/api/documents/delete-all')
def delete_all_documents():
    try:
        deleted_files = []
        failed_files = []
        
        # 1. Delete all files in the uploads directory
        upload_dir = os.path.join(os.path.dirname(__file__), '..', 'uploads')
        for filename in os.listdir(upload_dir):
            file_path = os.path.join(upload_dir, filename)
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
                    deleted_files.append(filename)
            except Exception as e:
                failed_files.append(f"{filename}: {str(e)}")
        
        # 2. Clear the ChromaDB collection
        result = ChromaDBHandler.clear_collection()
        if not result.get("success", False):
            return JSONResponse({
                "success": False, 
                "message": f"Failed to clear ChromaDB: {result.get('error', 'Unknown error')}",
                "deleted_files": deleted_files,
                "failed_files": failed_files
            }), 500
        
        # 3. Clear the uploaded_files table in the database
        db_path = os.path.join(os.path.dirname(__file__), '..', 'database', 'chat_history.db')
        
        # Use SQLite directly for this operation
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM uploaded_files")
        conn.commit()
        conn.close()
        
        # Log success details
        print(f"Document deletion completed. Files deleted: {len(deleted_files)}, Files failed: {len(failed_files)}")
        print(f"ChromaDB collection cleared: {result.get('message')}")
        
        return JSONResponse({
            "success": True, 
            "message": "All documents deleted successfully",
            "files_deleted": len(deleted_files),
            "vector_store_cleared": True
        })
    except Exception as e:
        print(f"Error in delete_all_documents: {str(e)}")
        return JSONResponse({"success": False, "message": str(e)}), 500
# Chat history API endpoints
@app.post("/api/conversations")
async def create_conversation(conversation: ConversationCreate):
    """Create a new conversation"""
    try:
        conversation_id = chat_history.create_conversation(
            title=conversation.title,
            metadata=conversation.metadata
        )
        return {"conversation_id": conversation_id, "title": conversation.title}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating conversation: {str(e)}")

@app.get("/api/conversations")
async def get_conversations(limit: int = 50, offset: int = 0):
    """Get a list of all conversations"""
    try:
        conversations = chat_history.get_all_conversations(limit=limit, offset=offset)
        return {"conversations": conversations}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error retrieving conversations: {str(e)}")

@app.get("/api/conversations/{conversation_id}")
async def get_conversation(conversation_id: str):
    """Get a specific with all its messages"""
    try:
        conversation = chat_history.get_conversation(conversation_id)
        return conversation
    except Exception as e:
        raise HTTPException(status_code=404, detail=f"Conversation not found: {str(e)}")

@app.put("/api/conversations/{conversation_id}")
async def update_conversation_title(conversation_id: str, update: ConversationUpdate):
    """Update a conversation's title"""
    try:
        success = chat_history.update_conversation_title(conversation_id, update.title)
        if not success:
            raise HTTPException(status_code=404, detail=f"Conversation {conversation_id} not found")
        return {"message": "updated successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating conversation: {str(e)}")

@app.delete("/api/conversations/{conversation_id}")
async def delete_conversation(conversation_id: str):
    """Delete a and all its messages"""
    try:
        success = chat_history.delete_conversation(conversation_id)
        if not success:
            raise HTTPException(status_code=404, detail=f"Conversation {conversation_id} not found")
        return {"message": "Conversation deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting conversation: {str(e)}")

@app.post("/api/conversations/{conversation_id}/messages")
async def add_message(conversation_id: str, message: MessageCreate):
    """Add a message to a conversation"""
    try:
        message_id = chat_history.add_message(
            conversation_id=conversation_id,
            role=message.role,
            content=message.content
        )
        return {"message_id": message_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error adding message: {str(e)}")
    
# DB routes for chat history to match the frontend's expected endpoints
@app.get("/db/get-all-chats")
async def get_all_chats():
    """Get all chats for the frontend"""
    print('inside get ALL chats---------------------')
    try:
        conversations = chat_history.get_all_conversations(limit=100)
        # Return in the format expected by frontend
        return conversations
    except Exception as e:
        print(f"Error getting all chats: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to retrieve chats: {str(e)}"}
        )

@app.get("/db/get-chat/{chat_id}")
async def get_chat(chat_id: str):
    """Get a specific chat by ID"""
    try:
        conversation = chat_history.get_conversation(chat_id)
        return conversation
    except Exception as e:
        print(f"Error getting chat {chat_id}: {str(e)}")
        return JSONResponse(
            status_code=404,
            content={"error": f"Chat not found: {str(e)}"}
        )

@app.post("/db/save-chat")
async def save_chat(request: Request):
    """Save a chat to the database"""
    print('-----------------> Inside SaveChat : ')
    try:
        data = await request.json()
        chat_id = data.get("id")
        title = data.get("title", "Untitled Chat")
        messages = data.get("messages", [])
        
        # Check if this is a new conversation or an update
        if not chat_id or chat_id == "null" or chat_id == "undefined":
            # Create new conversation
            metadata = {
                "system_prompt": data.get("systemPrompt", "You are a helpful assistant.")
            }
            chat_id = chat_history.create_conversation(title=title, metadata=metadata)
        else:
            # Update existing conversation title if provided
            try:
                if title:
                    chat_history.update_conversation_title(chat_id, title)
            
            # Get existing conversation to check if we need to add messages
                print('save-chat ----->'+chat_id)
                existing = []
                try:
                    existing = chat_history.get_conversation(chat_id)
                    print('Conversation FOUND!!!!!!!!!!!!!!!!')
                except:
                    print('no existing conversation !!!!!')
                
                print('----------existing---------')
                print(existing)
                print('----------existing---------')
                existing_msg_count = len(existing)#.get("messages", []))
                print(f'----------existing_msg_count--------- {existing_msg_count}')
                existing_messages = existing.get("messages", [])
                existing_set = set(
                    (m.get("role", ""), m.get("content", "").strip()) for m in existing_messages
                )
                # If all new messages are already in existing messages, skip
                incoming_set = set(
                    (m.get("role", ""), m.get("content", "").strip()) for m in messages
                )
                
                # If the number of messages is the same, no need to add more
                
                
                # if existing_msg_count >= len(messages):
                #     return {"id": chat_id, "success": True, "message": "Chat already up to date"}

                if incoming_set.issubset(existing_set):
                    print("!!!!!!!!!No new messages to add!!!!!")
                    return {"id": chat_id, "success": True, "message": "Chat already up to date"}

            except:
                # If the conversation doesn't exist, create it
                metadata = {
                    "system_prompt": data.get("systemPrompt", "You are a helpful assistant.")
                }
                print('adding new conversation ----->'+chat_id)
                chat_id = chat_history.create_conversation(title=title, metadata=metadata,conversation_id=chat_id)
                # chat_id = chat_history.create_conversation(title=title, metadata=metadata)
        
        # Add messages to the conversation
        # for msg in messages:
        #     # Skip adding messages that may already exist
        #     chat_history.add_message(
        #         conversation_id=chat_id,
        #         role=msg.get("role", "user"),
        #         content=msg.get("content", "")
        #     )
            
        # return {"id": chat_id, "success": True}

        # Add only new messages
        for msg in messages:
            role = msg.get("role", "user")
            content = msg.get("content", "").strip()
            if (role, content) not in existing_set:
                chat_history.add_message(
                    conversation_id=chat_id,
                    role=role,
                    content=content
                )

        return {"id": chat_id, "success": True}
    except Exception as e:
        print(f"Error saving chat: {str(e)}")
        return JSONResponse(
            status_code=500, 
            content={"error": f"Failed to save chat: {str(e)}"}
        )

@app.delete("/db/delete-chat/{chat_id}")
async def delete_chat(chat_id: str):
    """Delete a chat from the database"""
    try:
        success = chat_history.delete_conversation(chat_id)
        if success:
            return {"success": True, "message": "Chat deleted successfully"}
        else:
            return JSONResponse(
                status_code=404,
                content={"success": False, "message": "Chat not found"}
            )
    except Exception as e:
        print(f"Error deleting chat {chat_id}: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"success": False, "message": f"Failed to delete chat: {str(e)}"}
        )

# Document listing endpoint
@app.get("/api/documents")
async def get_documents():
    """Get a list of uploaded documents"""
    try:
        # Connect to the database
        conn = sqlite3.connect(chat_history.db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        # Get all unique files from the uploaded_files table
        cursor.execute("SELECT id, filename, filehash FROM uploaded_files ORDER BY id DESC")
        files = [dict(row) for row in cursor.fetchall()]
        conn.close()
        
        return {"files": files}
    except Exception as e:
        print(f"Error retrieving documents: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Failed to retrieve documents: {str(e)}"}
        )

# Alternative route for documents
@app.get("/documents")
async def documents_redirect():
    """Redirect to /api/documents for compatibility"""
    return await get_documents()
# Other endpoints remain unchanged...
# Goal Analyzer FastAPI Backend with Ollama
# Requirements: pip install fastapi uvicorn python-multipart aiofiles PyPDF2 python-docx pandas openpyxl xlrd python-magic httpx




# NEHA
# for goal analyzer


import os
import json
import tempfile
import magic
import asyncio
from datetime import datetime
from typing import List, Dict, Any, Optional
import logging
from pathlib import Path

from fastapi import FastAPI, File, UploadFile, Form, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import aiofiles
import httpx

# File processing libraries
import PyPDF2
from docx import Document
import pandas as pd
import csv
from io import BytesIO, StringIO

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Goal Analyzer API", version="1.0.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration
UPLOAD_FOLDER = tempfile.gettempdir()
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB
ALLOWED_EXTENSIONS = {
    'pdf', 'doc', 'docx', 'txt', 'json', 'csv', 'xlsx', 'xls'
}

# Ollama Configuration
OLLAMA_BASE_URL = "http://localhost:11434"
DEFAULT_MODEL = "llama3.2"  # Change to your preferred model

class FileProcessor:
    """Handles different file type processing"""
    
    @staticmethod
    async def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing PDF: {str(e)}")

    @staticmethod
    async def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing DOCX: {str(e)}")

    @staticmethod
    async def extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            return file_content.decode('utf-8', errors='ignore').strip()
        except Exception as e:
            logger.error(f"Error extracting TXT text: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing TXT: {str(e)}")

    @staticmethod
    async def extract_text_from_csv(file_content: bytes) -> str:
        """Extract text from CSV file"""
        try:
            csv_text = file_content.decode('utf-8', errors='ignore')
            csv_file = StringIO(csv_text)
            reader = csv.reader(csv_file)
            
            text = ""
            for row in reader:
                text += " | ".join(row) + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting CSV text: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing CSV: {str(e)}")

    @staticmethod
    async def extract_text_from_excel(file_content: bytes, filename: str) -> str:
        """Extract text from Excel file"""
        try:
            excel_file = BytesIO(file_content)
            
            # Try to read as xlsx first, then xls
            try:
                if filename.endswith('.xlsx'):
                    df_dict = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
                else:
                    df_dict = pd.read_excel(excel_file, sheet_name=None, engine='xlrd')
            except:
                # Fallback to auto-detection
                df_dict = pd.read_excel(excel_file, sheet_name=None)
            
            text = ""
            for sheet_name, df in df_dict.items():
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string(index=False) + "\n\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting Excel text: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing Excel: {str(e)}")

    @staticmethod
    async def extract_text_from_json(file_content: bytes) -> str:
        """Extract text from JSON file"""
        try:
            json_text = file_content.decode('utf-8', errors='ignore')
            json_data = json.loads(json_text)
            
            # Convert JSON to readable text format
            text = json.dumps(json_data, indent=2, ensure_ascii=False)
            return text
        except Exception as e:
            logger.error(f"Error extracting JSON text: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Error processing JSON: {str(e)}")

class OllamaClient:
    """Client for interacting with Ollama API"""
    
    def __init__(self, base_url: str = OLLAMA_BASE_URL):
        self.base_url = base_url
        self.client = httpx.AsyncClient(timeout=120.0)
    
    async def generate_response(self, prompt: str, model: str = DEFAULT_MODEL) -> str:
        """Generate response using Ollama"""
        try:
            response = await self.client.post(
                f"{self.base_url}/api/generate",
                json={
                    "model": model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.7,
                        "top_p": 0.9,
                        "max_tokens": 2000
                    }
                }
            )
            response.raise_for_status()
            result = response.json()
            return result.get("response", "")
        except httpx.RequestError as e:
            logger.error(f"Request error with Ollama: {str(e)}")
            raise HTTPException(status_code=503, detail="AI service unavailable")
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            raise HTTPException(status_code=500, detail=f"AI processing error: {str(e)}")
    
    async def check_model_availability(self, model: str = DEFAULT_MODEL) -> bool:
        """Check if the specified model is available"""
        try:
            response = await self.client.get(f"{self.base_url}/api/tags")
            response.raise_for_status()
            models = response.json()
            model_names = [m["name"] for m in models.get("models", [])]
            return model in model_names or f"{model}:latest" in model_names
        except:
            return False

class GoalAnalyzer:
    """Main goal analysis logic"""
    
    def __init__(self):
        self.file_processor = FileProcessor()
        self.ollama_client = OllamaClient()
    
    def get_file_extension(self, filename: str) -> str:
        """Get file extension"""
        return filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    def is_allowed_file(self, filename: str) -> bool:
        """Check if file type is allowed"""
        return self.get_file_extension(filename) in ALLOWED_EXTENSIONS
    
    async def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """Extract text based on file type"""
        extension = self.get_file_extension(filename)
        
        if extension == 'pdf':
            return await self.file_processor.extract_text_from_pdf(file_content)
        elif extension in ['doc', 'docx']:
            return await self.file_processor.extract_text_from_docx(file_content)
        elif extension == 'txt':
            return await self.file_processor.extract_text_from_txt(file_content)
        elif extension == 'csv':
            return await self.file_processor.extract_text_from_csv(file_content)
        elif extension in ['xlsx', 'xls']:
            return await self.file_processor.extract_text_from_excel(file_content, filename)
        elif extension == 'json':
            return await self.file_processor.extract_text_from_json(file_content)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {extension}")
    
    def create_goal_analysis_prompt(self, texts: List[Dict[str, str]]) -> str:
        """Create a comprehensive prompt for goal analysis"""
        
        combined_content = ""
        for item in texts:
            combined_content += f"\n--- File: {item['filename']} ---\n{item['content']}\n"
        
        prompt = f"""
You are an expert goal analyst and strategic planning consultant. Analyze the following documents and provide a comprehensive goal analysis report.

DOCUMENTS TO ANALYZE:
{combined_content}

Please provide a detailed analysis covering these areas:

1. **IDENTIFIED GOALS AND OBJECTIVES**
   - List all explicit and implicit goals mentioned
   - Categorize them (short-term, medium-term, long-term)
   - Identify primary vs secondary objectives

2. **GOAL CLARITY AND SPECIFICITY**
   - Assess how well-defined each goal is
   - Identify SMART criteria compliance (Specific, Measurable, Achievable, Relevant, Time-bound)
   - Highlight vague or unclear objectives

3. **PRIORITY ANALYSIS**
   - Rank goals by apparent importance/urgency
   - Identify potential conflicts between goals
   - Suggest priority ordering

4. **FEASIBILITY ASSESSMENT**
   - Evaluate realistic achievability of each goal
   - Identify resource requirements mentioned
   - Highlight potential obstacles or challenges

5. **STRATEGIC ALIGNMENT**
   - Assess how goals align with overall vision/mission (if present)
   - Identify synergies between different objectives
   - Spot potential contradictions

6. **ACTION PLANNING GAPS**
   - Identify goals lacking clear action steps
   - Suggest missing milestones or checkpoints
   - Recommend implementation strategies

7. **RECOMMENDATIONS**
   - Prioritized action items
   - Goal refinement suggestions
   - Strategic improvements
   - Next steps for goal achievement

Format your response in clear sections with bullet points and actionable insights. Be specific and provide practical recommendations.
"""
        return prompt
    
    async def analyze_goals(self, file_contents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Perform comprehensive goal analysis"""
        try:
            # Extract text from all files
            extracted_texts = []
            for file_data in file_contents:
                text = await self.extract_text_from_file(
                    file_data['content'], 
                    file_data['filename']
                )
                extracted_texts.append({
                    'filename': file_data['filename'],
                    'content': text[:5000]  # Limit content length
                })
            
            # Check if Ollama model is available
            model_available = await self.ollama_client.check_model_availability()
            if not model_available:
                raise HTTPException(
                    status_code=503, 
                    detail=f"AI model '{DEFAULT_MODEL}' not available. Please ensure Ollama is running and the model is installed."
                )
            
            # Create analysis prompt
            prompt = self.create_goal_analysis_prompt(extracted_texts)
            
            # Generate analysis using Ollama
            analysis_result = await self.ollama_client.generate_response(prompt)
            
            return {
                'success': True,
                'analysis': analysis_result,
                'files_processed': len(file_contents),
                'timestamp': datetime.now().isoformat(),
                'model_used': DEFAULT_MODEL
            }
            
        except Exception as e:
            logger.error(f"Goal analysis error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

# Initialize goal analyzer
goal_analyzer = GoalAnalyzer()

@app.get("/")
async def root():
    """Root endpoint"""
    return {"message": "Goal Analyzer API", "version": "1.0.0"}

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    try:
        # Check Ollama availability
        model_available = await goal_analyzer.ollama_client.check_model_availability()
        return {
            "status": "healthy",
            "ollama_available": model_available,
            "model": DEFAULT_MODEL,
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.post("/api/analyze-goals")
async def analyze_goals_endpoint(
    files: List[UploadFile] = File(...),
    analysis_type: str = Form(default="goal_analysis")
):
    """Main endpoint for goal analysis"""
    try:
        if not files:
            raise HTTPException(status_code=400, detail="No files uploaded")
        
        # Validate files
        file_contents = []
        for file in files:
            # Check file size
            if file.size > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=400, 
                    detail=f"File {file.filename} exceeds maximum size of {MAX_FILE_SIZE/1024/1024}MB"
                )
            
            # Check file type
            if not goal_analyzer.is_allowed_file(file.filename):
                raise HTTPException(
                    status_code=400, 
                    detail=f"File type not supported: {file.filename}"
                )
            
            # Read file content
            content = await file.read()
            file_contents.append({
                'filename': file.filename,
                'content': content,
                'size': len(content)
            })
        
        # Perform analysis
        result = await goal_analyzer.analyze_goals(file_contents)
        
        return JSONResponse(content=result)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in analyze_goals_endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.get("/api/models")
async def get_available_models():
    """Get available Ollama models"""
    try:
        response = await goal_analyzer.ollama_client.client.get(f"{OLLAMA_BASE_URL}/api/tags")
        response.raise_for_status()
        models = response.json()
        return {
            "models": [m["name"] for m in models.get("models", [])],
            "current_model": DEFAULT_MODEL
        }
    except Exception as e:
        raise HTTPException(status_code=503, detail="Could not fetch models from Ollama")

@app.post("/api/pull-model")
async def pull_model(model_name: str = Form(...)):
    """Pull a new model in Ollama"""
    try:
        response = await goal_analyzer.ollama_client.client.post(
            f"{OLLAMA_BASE_URL}/api/pull",
            json={"name": model_name}
        )
        response.raise_for_status()
        return {"message": f"Model {model_name} pull initiated"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to pull model: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
